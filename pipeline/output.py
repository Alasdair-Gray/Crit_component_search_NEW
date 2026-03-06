"""
Stage 5 – Report Generation
==============================
Renders a formatted .docx certification report from the compiled pipeline
output.

Responsibilities
----------------
- Accept a ``PipelineOutput`` object and an output path.
- Generate a formatted .docx report using python-docx:
    - Landscape A4 with 1.5 cm margins.
    - Title, subtitle, and generation date.
    - A single certification table for all assemblies, with dark-blue header
      row, light-blue section divider rows, and alternating data row shading.
    - Clickable hyperlinks in the Source + URL column.
    - Thin grey cell borders throughout.
    - Page-number footer and project-name header.
    - Italicised footer note listing any components that need manual review.
- Write the file to the specified output path.
- Return the resolved output path.

API contract
------------
Input  : ``PipelineOutput``, ``output_path: str | Path``
Output : ``Path`` (resolved path to the written file)

Example usage
-------------
::

    from pipeline.output import generate_report

    report_path = generate_report(pipeline_output, "reports/acme_widget_v2.docx")
    print(f"Report written to {report_path}")
"""

from __future__ import annotations

import logging
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from pipeline.models import PipelineOutput

log = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Colour constants
# ---------------------------------------------------------------------------

_DARK_BLUE_HEX = "1F4E79"    # title text, table header background
_SECTION_BLUE_HEX = "D6E4F0" # assembly section divider background
_ROW_ALT_HEX = "F0F5FF"      # alternating data row shading
_ROW_WHITE_HEX = "FFFFFF"    # plain data row shading
_BORDER_GREY_HEX = "AAAAAA"  # cell border colour

_DARK_BLUE_RGB = RGBColor(0x1F, 0x4E, 0x79)
_WHITE_RGB = RGBColor(0xFF, 0xFF, 0xFF)
_GREY_RGB = RGBColor(0x80, 0x80, 0x80)

# ---------------------------------------------------------------------------
# Table layout
# ---------------------------------------------------------------------------

_COL_HEADERS = ["Component", "Standards", "Certificates", "Source + URL"]
_COL_WIDTHS = [Cm(5.0), Cm(7.5), Cm(7.0), Cm(7.2)]  # total 26.7 cm
_TOTAL_TABLE_WIDTH = Cm(26.7)

# ---------------------------------------------------------------------------
# XML helper utilities
# ---------------------------------------------------------------------------


def _shade_cell(cell, fill_hex: str) -> None:
    """Set the background fill colour of *cell*."""
    tcPr = cell._tc.get_or_add_tcPr()
    for existing in tcPr.findall(qn("w:shd")):
        tcPr.remove(existing)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def _set_cell_borders(cell, color: str = _BORDER_GREY_HEX) -> None:
    """Apply thin single-line borders on all four sides of *cell*."""
    tcPr = cell._tc.get_or_add_tcPr()
    for existing in tcPr.findall(qn("w:tcBorders")):
        tcPr.remove(existing)
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")       # 0.5 pt line
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _add_hyperlink(paragraph, text: str, url: str) -> None:
    """Insert a clickable hyperlink run into *paragraph*."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run_el = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)
    run_el.append(rPr)

    t_el = OxmlElement("w:t")
    t_el.text = text
    run_el.append(t_el)

    hyperlink.append(run_el)
    paragraph._p.append(hyperlink)


def _add_page_number_field(paragraph) -> None:
    """Insert a ``PAGE`` field into *paragraph* for automatic page numbering."""
    # begin
    run = paragraph.add_run()
    fldChar = OxmlElement("w:fldChar")
    fldChar.set(qn("w:fldCharType"), "begin")
    run._r.append(fldChar)
    # instruction
    run = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.text = "PAGE"
    run._r.append(instr)
    # end
    run = paragraph.add_run()
    fldChar = OxmlElement("w:fldChar")
    fldChar.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar)


# ---------------------------------------------------------------------------
# Document setup
# ---------------------------------------------------------------------------


def _setup_document(
    doc: Document,
    project_name: str,
    source_document: str,
) -> None:
    """Configure landscape A4 page, margins, page header, and page footer."""
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    for attr in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
        setattr(section, attr, Cm(1.5))

    # Page header – right-aligned project name and source document
    header_para = section.header.paragraphs[0]
    header_para.clear()
    run = header_para.add_run(f"{project_name}  |  {source_document}")
    run.font.size = Pt(9)
    run.font.color.rgb = _GREY_RGB
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Page footer – centred page number
    footer_para = section.footer.paragraphs[0]
    footer_para.clear()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_page_number_field(footer_para)


# ---------------------------------------------------------------------------
# Title block
# ---------------------------------------------------------------------------


def _write_title(
    doc: Document,
    project_name: str,
    source_document: str,
    generated_at,
) -> None:
    """Write the report title, subtitle, and generation timestamp."""
    title_para = doc.add_paragraph()
    run = title_para.add_run("Component Certification Register")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = _DARK_BLUE_RGB

    sub_para = doc.add_paragraph()
    run = sub_para.add_run(
        f"Project: {project_name}  |  Source: {source_document}  |  "
        f"Generated: {generated_at.strftime('%Y-%m-%d %H:%M UTC')}"
    )
    run.font.size = Pt(11)
    run.font.color.rgb = _GREY_RGB


# ---------------------------------------------------------------------------
# Table construction
# ---------------------------------------------------------------------------


def _format_header_row(table) -> None:
    """Format the first (header) row with dark-blue background and white bold text."""
    row = table.rows[0]
    for cell, heading, width in zip(row.cells, _COL_HEADERS, _COL_WIDTHS):
        cell.width = width
        _shade_cell(cell, _DARK_BLUE_HEX)
        _set_cell_borders(cell)
        para = cell.paragraphs[0]
        para.clear()
        run = para.add_run(heading)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = _WHITE_RGB


def _add_section_divider(table, assembly_name: str) -> None:
    """Append a light-blue section divider row spanning all four columns."""
    row = table.add_row()
    merged = row.cells[0].merge(row.cells[3])
    merged.width = _TOTAL_TABLE_WIDTH
    _shade_cell(merged, _SECTION_BLUE_HEX)
    _set_cell_borders(merged)
    para = merged.paragraphs[0]
    para.clear()
    run = para.add_run(assembly_name)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = _DARK_BLUE_RGB


def _add_data_row(
    table,
    row_index: int,
    component_name: str,
    standards_text: str,
    certificates_text: str,
    source_entries: list[tuple[str, str]],
) -> None:
    """Append one data row (one row per component) with alternating shading.

    Parameters
    ----------
    source_entries:
        List of ``(url, display_name)`` pairs — one hyperlink is written per
        unique source URL.  An empty list results in an em-dash cell.
    """
    fill = _ROW_ALT_HEX if row_index % 2 == 1 else _ROW_WHITE_HEX
    row = table.add_row()
    cells = row.cells

    # Columns 0–2: plain text (component, standards, certificates)
    for cell, text, width in zip(
        cells[:3],
        [component_name, standards_text, certificates_text],
        _COL_WIDTHS[:3],
    ):
        cell.width = width
        _shade_cell(cell, fill)
        _set_cell_borders(cell)
        para = cell.paragraphs[0]
        para.clear()
        run = para.add_run(text)
        run.font.size = Pt(9)

    # Column 3: one hyperlink per unique source URL, stacked in the cell
    source_cell = cells[3]
    source_cell.width = _COL_WIDTHS[3]
    _shade_cell(source_cell, fill)
    _set_cell_borders(source_cell)

    if source_entries:
        first_para = source_cell.paragraphs[0]
        first_para.clear()
        for idx, (url, name) in enumerate(source_entries):
            # Add a new paragraph for every source after the first
            para = first_para if idx == 0 else source_cell.add_paragraph()
            if url:
                _add_hyperlink(para, name or url, url)
            else:
                run = para.add_run(name or "–")
                run.font.size = Pt(9)
    else:
        source_para = source_cell.paragraphs[0]
        source_para.clear()
        run = source_para.add_run("–")
        run.font.size = Pt(9)


def _write_table(doc: Document, data: PipelineOutput) -> None:
    """Build the complete certification table for all assemblies.

    Each component occupies **one row**.  The Standards column lists all
    technical specifications found; the Certificates column lists all issued
    compliance document numbers.  Multiple source URLs are stacked in the
    Source + URL cell as individual hyperlinks.
    """
    table = doc.add_table(rows=1, cols=4)
    _format_header_row(table)

    data_row_index = 0
    for assembly_name, results in data.results_by_assembly.items():
        _add_section_divider(table, assembly_name)
        for result in results:
            # Split CertificationFound items by kind
            standards = [c for c in result.certifications if c.kind == "standard"]
            certificates = [c for c in result.certifications if c.kind == "certificate"]

            # Build Standards cell text (one standard per line)
            if standards:
                standards_text = "\n".join(c.standard for c in standards)
            else:
                standards_text = "–"

            # Build Certificates cell text (cert number + standard reference per line)
            if certificates:
                cert_lines: list[str] = []
                for c in certificates:
                    if c.cert_number and c.standard:
                        cert_lines.append(f"{c.cert_number}  ({c.standard})")
                    elif c.cert_number:
                        cert_lines.append(c.cert_number)
                    else:
                        cert_lines.append(c.standard)
                certificates_text = "\n".join(cert_lines)
            else:
                certificates_text = "–"

            # Collect unique source entries (preserving order)
            seen_urls: set[str] = set()
            source_entries: list[tuple[str, str]] = []
            for c in result.certifications:
                if c.source_url and c.source_url not in seen_urls:
                    seen_urls.add(c.source_url)
                    source_entries.append((c.source_url, c.source_name))

            _add_data_row(
                table,
                data_row_index,
                result.enriched_component.name,
                standards_text,
                certificates_text,
                source_entries,
            )
            data_row_index += 1


# ---------------------------------------------------------------------------
# Footer note
# ---------------------------------------------------------------------------


def _write_footer_note(doc: Document, components_needing_review: list[str]) -> None:
    """Append an italicised note listing components flagged for manual review."""
    if not components_needing_review:
        return
    doc.add_paragraph()  # spacer
    note_para = doc.add_paragraph()
    label = note_para.add_run("Components requiring manual review: ")
    label.bold = True
    label.italic = True
    label.font.size = Pt(9)
    label.font.color.rgb = _GREY_RGB
    items = note_para.add_run("; ".join(components_needing_review))
    items.italic = True
    items.font.size = Pt(9)
    items.font.color.rgb = _GREY_RGB


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------


def generate_report(pipeline_output: PipelineOutput, output_path: str | Path) -> Path:
    """Render *pipeline_output* as a formatted ``.docx`` file.

    Parameters
    ----------
    pipeline_output:
        The structured result from :mod:`pipeline.compile`.
    output_path:
        Destination path for the generated report.

    Returns
    -------
    Path
        Resolved, absolute path to the written ``.docx`` file.

    Raises
    ------
    OSError
        If the output directory does not exist or is not writable.
    """
    output_path = Path(output_path).resolve()

    doc = Document()
    _setup_document(doc, pipeline_output.project_name, pipeline_output.source_document)
    _write_title(
        doc,
        pipeline_output.project_name,
        pipeline_output.source_document,
        pipeline_output.generated_at,
    )
    _write_table(doc, pipeline_output)
    _write_footer_note(doc, pipeline_output.components_needing_review)

    doc.save(str(output_path))
    log.info("Report saved: %s", output_path)
    return output_path
