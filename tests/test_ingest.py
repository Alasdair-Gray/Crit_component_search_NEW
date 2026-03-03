"""
Unit tests for pipeline/ingest.py.

The LLM is always mocked so no real API calls are made.
"""

from __future__ import annotations

import json
from pathlib import Path
from unittest.mock import MagicMock

import pytest
from docx import Document


# ---------------------------------------------------------------------------
# Fixtures – small .docx files created in memory
# ---------------------------------------------------------------------------


def _make_component_table_docx(tmp_path: Path) -> Path:
    """A .docx with a Heading 1 and a 3-column component table beneath it."""
    doc = Document()
    doc.add_heading("Power Assembly", level=1)

    table = doc.add_table(rows=3, cols=3)
    headers = table.rows[0].cells
    headers[0].text = "Component"
    headers[1].text = "Part Number"
    headers[2].text = "Manufacturer"

    row1 = table.rows[1].cells
    row1[0].text = "DIN Rail PSU"
    row1[1].text = "HDR-100-12"
    row1[2].text = "Mean Well"

    row2 = table.rows[2].cells
    row2[0].text = "IEC Inlet"
    row2[1].text = "6100-42"
    row2[2].text = "Schurter"

    path = tmp_path / "test_spec.docx"
    doc.save(str(path))
    return path


def _make_empty_docx(tmp_path: Path) -> Path:
    """A .docx with no content at all."""
    doc = Document()
    path = tmp_path / "empty.docx"
    doc.save(str(path))
    return path


def _make_text_only_docx(tmp_path: Path) -> Path:
    """A .docx with prose but no components."""
    doc = Document()
    doc.add_paragraph("This document contains general project notes.")
    doc.add_paragraph("No component tables or lists are present.")
    path = tmp_path / "no_components.docx"
    doc.save(str(path))
    return path


def _make_multiassembly_docx(tmp_path: Path) -> Path:
    """A .docx with two assembly sections each containing a list."""
    doc = Document()

    doc.add_heading("Mains Cord", level=1)
    doc.add_paragraph("H07V2-K BASEC 3G1.5mm² mains flex")

    doc.add_heading("Socket/USB Module", level=1)
    doc.add_paragraph("WAGO 2002-1401 terminal block")
    doc.add_paragraph("Schurter 6100-42 IEC inlet")

    path = tmp_path / "multi_assembly.docx"
    doc.save(str(path))
    return path


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _mock_llm(components_payload: list[dict]) -> MagicMock:
    """Return a mock LLM that responds with the given components JSON."""
    mock = MagicMock()
    mock.complete.return_value = json.dumps({"components": components_payload})
    return mock


# ---------------------------------------------------------------------------
# Tests
# ---------------------------------------------------------------------------


class TestExtractComponents:
    def test_components_from_table(self, tmp_path):
        """Components in a table are extracted and mapped to Component objects."""
        from pipeline.ingest import extract_components

        docx_path = _make_component_table_docx(tmp_path)
        llm = _mock_llm(
            [
                {
                    "name": "Mean Well HDR-100-12",
                    "assembly": "Power Assembly",
                    "raw_text": "DIN Rail PSU | HDR-100-12 | Mean Well",
                    "part_number": "HDR-100-12",
                    "manufacturer": "Mean Well",
                    "description": "DIN rail power supply 100 W 12 V",
                },
                {
                    "name": "Schurter 6100-42",
                    "assembly": "Power Assembly",
                    "raw_text": "IEC Inlet | 6100-42 | Schurter",
                    "part_number": "6100-42",
                    "manufacturer": "Schurter",
                    "description": "IEC C14 inlet",
                },
            ]
        )

        components = extract_components(docx_path, llm=llm)

        assert len(components) == 2

        psu = components[0]
        assert psu.name == "Mean Well HDR-100-12"
        assert psu.part_number == "HDR-100-12"
        assert psu.manufacturer == "Mean Well"
        assert psu.assembly == "Power Assembly"
        assert "HDR-100-12" in psu.raw_text

        inlet = components[1]
        assert inlet.name == "Schurter 6100-42"
        assert inlet.part_number == "6100-42"
        assert inlet.manufacturer == "Schurter"

        # LLM should have been called exactly once
        llm.complete.assert_called_once()

    def test_empty_document_returns_empty_list_without_llm_call(self, tmp_path):
        """An empty .docx returns [] and never calls the LLM."""
        from pipeline.ingest import extract_components

        docx_path = _make_empty_docx(tmp_path)
        llm = MagicMock()

        components = extract_components(docx_path, llm=llm)

        assert components == []
        llm.complete.assert_not_called()

    def test_no_recognisable_components_returns_empty_list(self, tmp_path):
        """A document with prose but no components returns [] without crashing."""
        from pipeline.ingest import extract_components

        docx_path = _make_text_only_docx(tmp_path)
        llm = _mock_llm([])

        components = extract_components(docx_path, llm=llm)

        assert components == []
        llm.complete.assert_called_once()

    def test_optional_fields_can_be_null(self, tmp_path):
        """Components with null part_number / manufacturer are accepted."""
        from pipeline.ingest import extract_components

        docx_path = _make_text_only_docx(tmp_path)
        llm = _mock_llm(
            [
                {
                    "name": "E254552 AWM 1015",
                    "assembly": "Mains Cord",
                    "raw_text": "E254552 AWM 1015",
                    "part_number": None,
                    "manufacturer": None,
                    "description": "Wire marking / cable designation",
                }
            ]
        )

        components = extract_components(docx_path, llm=llm)

        assert len(components) == 1
        c = components[0]
        assert c.name == "E254552 AWM 1015"
        assert c.part_number is None
        assert c.manufacturer is None

    def test_multiple_assemblies(self, tmp_path):
        """Components from different assemblies are all extracted."""
        from pipeline.ingest import extract_components

        docx_path = _make_multiassembly_docx(tmp_path)
        llm = _mock_llm(
            [
                {
                    "name": "H07V2-K BASEC 3G1.5mm²",
                    "assembly": "Mains Cord",
                    "raw_text": "H07V2-K BASEC 3G1.5mm² mains flex",
                    "part_number": None,
                    "manufacturer": None,
                    "description": "Mains flex cable",
                },
                {
                    "name": "WAGO 2002-1401",
                    "assembly": "Socket/USB Module",
                    "raw_text": "WAGO 2002-1401 terminal block",
                    "part_number": "2002-1401",
                    "manufacturer": "WAGO",
                    "description": "Terminal block",
                },
                {
                    "name": "Schurter 6100-42",
                    "assembly": "Socket/USB Module",
                    "raw_text": "Schurter 6100-42 IEC inlet",
                    "part_number": "6100-42",
                    "manufacturer": "Schurter",
                    "description": "IEC C14 inlet",
                },
            ]
        )

        components = extract_components(docx_path, llm=llm)

        assert len(components) == 3
        assemblies = {c.assembly for c in components}
        assert assemblies == {"Mains Cord", "Socket/USB Module"}

    def test_file_not_found_raises(self, tmp_path):
        """FileNotFoundError is raised for a missing file."""
        from pipeline.ingest import extract_components

        with pytest.raises(FileNotFoundError):
            extract_components(tmp_path / "does_not_exist.docx")

    def test_invalid_llm_json_raises_value_error(self, tmp_path):
        """A non-JSON LLM response raises ValueError."""
        from pipeline.ingest import extract_components

        docx_path = _make_text_only_docx(tmp_path)
        llm = MagicMock()
        llm.complete.return_value = "Sorry, I could not parse the document."

        with pytest.raises(ValueError, match="unparseable response"):
            extract_components(docx_path, llm=llm)

    def test_malformed_component_entry_is_skipped(self, tmp_path):
        """A component dict missing 'name' is skipped; others still returned."""
        from pipeline.ingest import extract_components

        docx_path = _make_text_only_docx(tmp_path)
        llm = _mock_llm(
            [
                # missing required 'name' key
                {
                    "assembly": "Power Assembly",
                    "raw_text": "bad entry",
                    "part_number": None,
                    "manufacturer": None,
                    "description": None,
                },
                {
                    "name": "Good Component",
                    "assembly": "Power Assembly",
                    "raw_text": "Good Component text",
                    "part_number": "ABC-123",
                    "manufacturer": "Acme",
                    "description": None,
                },
            ]
        )

        components = extract_components(docx_path, llm=llm)

        assert len(components) == 1
        assert components[0].name == "Good Component"

    def test_llm_response_wrapped_in_markdown_fence(self, tmp_path):
        """JSON wrapped in ```json ... ``` fences is parsed correctly."""
        from pipeline.ingest import extract_components

        docx_path = _make_text_only_docx(tmp_path)
        llm = MagicMock()
        llm.complete.return_value = (
            "```json\n"
            + json.dumps(
                {
                    "components": [
                        {
                            "name": "Fused Component",
                            "assembly": "Main",
                            "raw_text": "Fused Component",
                            "part_number": "F-1",
                            "manufacturer": "Brand",
                            "description": None,
                        }
                    ]
                }
            )
            + "\n```"
        )

        components = extract_components(docx_path, llm=llm)

        assert len(components) == 1
        assert components[0].name == "Fused Component"
