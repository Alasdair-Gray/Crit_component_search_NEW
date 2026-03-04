"""
Unit tests for pipeline/output.py.

Tests verify the .docx is structurally correct without relying on visual
rendering.  All tests write to a pytest tmp_path directory.
"""

from __future__ import annotations

from datetime import datetime
from pathlib import Path

import pytest

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn

from pipeline.models import (
    CertificationFound,
    CertificationResult,
    Confidence,
    EnrichedComponent,
    PipelineOutput,
)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _make_enriched(
    name: str,
    assembly: str = "Main Assembly",
    manufacturer_uncertain: bool = False,
) -> EnrichedComponent:
    return EnrichedComponent(
        name=name,
        assembly=assembly,
        raw_text=name,
        part_number="P-001",
        manufacturer="Acme",
        confirmed_manufacturer="Acme",
        standardised_part_number="P-001",
        component_type="PSU",
        search_queries=["q1"],
        manufacturer_uncertain=manufacturer_uncertain,
    )


def _make_cert(
    standard: str = "UL 508",
    source_url: str = "https://ul.com/listing/E171376",
    source_name: str = "ul.com",
    cert_number: str | None = "E171376",
) -> CertificationFound:
    return CertificationFound(
        standard=standard,
        cert_number=cert_number,
        scope="Safety certification for power supplies",
        source_url=source_url,
        source_name=source_name,
    )


def _make_result(
    component: EnrichedComponent | None = None,
    certifications: list[CertificationFound] | None = None,
    confidence: Confidence = Confidence.HIGH,
) -> CertificationResult:
    return CertificationResult(
        enriched_component=component or _make_enriched("Test Component"),
        certifications=certifications or [],
        confidence=confidence,
        search_log=[],
    )


def _make_pipeline_output(
    results: list[CertificationResult] | None = None,
    components_needing_review: list[str] | None = None,
    project_name: str = "Test Project",
    source_document: str = "spec.docx",
) -> PipelineOutput:
    results = results or [_make_result(certifications=[_make_cert()])]
    results_by_assembly: dict[str, list[CertificationResult]] = {}
    for r in results:
        assembly = r.enriched_component.assembly
        results_by_assembly.setdefault(assembly, []).append(r)
    return PipelineOutput(
        project_name=project_name,
        source_document=source_document,
        results_by_assembly=results_by_assembly,
        components_needing_review=components_needing_review or [],
        generated_at=datetime(2024, 6, 15, 9, 30, 0),
    )


def _generate(tmp_path: Path, data: PipelineOutput | None = None) -> Path:
    """Write a report to tmp_path and return the output path."""
    from pipeline.output import generate_report

    out = tmp_path / "report.docx"
    data = data or _make_pipeline_output()
    generate_report(data, out)
    return out


# ---------------------------------------------------------------------------
# File output
# ---------------------------------------------------------------------------


class TestFileOutput:
    def test_file_is_written(self, tmp_path):
        """generate_report creates a file at the specified path."""
        out = _generate(tmp_path)
        assert out.exists()
        assert out.suffix == ".docx"

    def test_returns_resolved_path(self, tmp_path):
        """Return value is a resolved Path object pointing to the output file."""
        from pipeline.output import generate_report

        out = tmp_path / "report.docx"
        result = generate_report(_make_pipeline_output(), out)
        assert isinstance(result, Path)
        assert result.is_absolute()
        assert result == out.resolve()

    def test_output_is_valid_docx(self, tmp_path):
        """The generated file can be reopened by python-docx without error."""
        out = _generate(tmp_path)
        doc = Document(str(out))  # would raise on corrupt file
        assert doc is not None


# ---------------------------------------------------------------------------
# Page setup
# ---------------------------------------------------------------------------


class TestPageSetup:
    def test_orientation_is_landscape(self, tmp_path):
        """The report section is set to landscape orientation."""
        out = _generate(tmp_path)
        doc = Document(str(out))
        assert doc.sections[0].orientation == WD_ORIENT.LANDSCAPE

    def test_page_width_is_a4_landscape(self, tmp_path):
        """Page width is 29.7 cm (A4 landscape long edge)."""
        out = _generate(tmp_path)
        doc = Document(str(out))
        assert abs(doc.sections[0].page_width.cm - 29.7) < 0.1

    def test_page_height_is_a4_landscape(self, tmp_path):
        """Page height is 21.0 cm (A4 landscape short edge)."""
        out = _generate(tmp_path)
        doc = Document(str(out))
        assert abs(doc.sections[0].page_height.cm - 21.0) < 0.1

    def test_margins_are_1_5_cm(self, tmp_path):
        """All four margins are set to 1.5 cm."""
        out = _generate(tmp_path)
        doc = Document(str(out))
        section = doc.sections[0]
        for attr in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
            assert abs(getattr(section, attr).cm - 1.5) < 0.05, (
                f"{attr} expected 1.5 cm, got {getattr(section, attr).cm}"
            )


# ---------------------------------------------------------------------------
# Title and subtitle
# ---------------------------------------------------------------------------


class TestTitleBlock:
    def test_title_text_present(self, tmp_path):
        """'Component Certification Register' appears in the document body."""
        out = _generate(tmp_path)
        doc = Document(str(out))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Component Certification Register" in full_text

    def test_subtitle_contains_project_name(self, tmp_path):
        """The subtitle paragraph includes the project name."""
        out = _generate(tmp_path, _make_pipeline_output(project_name="Acme Widget v2"))
        doc = Document(str(out))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Acme Widget v2" in full_text

    def test_subtitle_contains_source_document(self, tmp_path):
        """The subtitle paragraph includes the source document filename."""
        out = _generate(tmp_path, _make_pipeline_output(source_document="acme_spec.docx"))
        doc = Document(str(out))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "acme_spec.docx" in full_text


# ---------------------------------------------------------------------------
# Table structure
# ---------------------------------------------------------------------------


class TestTableStructure:
    def test_table_has_four_columns(self, tmp_path):
        """The certification table is created with exactly four columns."""
        out = _generate(tmp_path)
        doc = Document(str(out))
        assert len(doc.tables) == 1
        assert len(doc.tables[0].columns) == 4

    def test_header_row_text(self, tmp_path):
        """The first table row contains the four expected column headings."""
        out = _generate(tmp_path)
        doc = Document(str(out))
        header_cells = doc.tables[0].rows[0].cells
        texts = [c.text for c in header_cells]
        assert texts == ["Component", "Cert / Standard", "Scope", "Source + URL"]

    def test_component_name_in_table(self, tmp_path):
        """The component name appears somewhere in the table rows."""
        component = _make_enriched("Mean Well HDR-100-12")
        result = _make_result(component, [_make_cert()])
        out = _generate(tmp_path, _make_pipeline_output([result]))
        doc = Document(str(out))
        all_cell_text = " ".join(
            c.text for row in doc.tables[0].rows for c in row.cells
        )
        assert "Mean Well HDR-100-12" in all_cell_text

    def test_certification_standard_in_table(self, tmp_path):
        """The certification standard appears in a table cell."""
        result = _make_result(certifications=[_make_cert("IEC 60320-1")])
        out = _generate(tmp_path, _make_pipeline_output([result]))
        doc = Document(str(out))
        all_cell_text = " ".join(
            c.text for row in doc.tables[0].rows for c in row.cells
        )
        assert "IEC 60320-1" in all_cell_text

    def test_no_certs_shows_dash(self, tmp_path):
        """A component with no certifications gets an em-dash in the standard column."""
        result = _make_result(certifications=[], confidence=Confidence.NOT_FOUND)
        out = _generate(tmp_path, _make_pipeline_output([result]))
        doc = Document(str(out))
        all_cell_text = " ".join(
            c.text for row in doc.tables[0].rows for c in row.cells
        )
        assert "–" in all_cell_text

    def test_assembly_name_appears_as_section_divider(self, tmp_path):
        """The assembly name appears in a merged section divider row."""
        component = _make_enriched("Widget", assembly="Power Board")
        result = _make_result(component, [_make_cert()])
        out = _generate(tmp_path, _make_pipeline_output([result]))
        doc = Document(str(out))
        all_cell_text = " ".join(
            c.text for row in doc.tables[0].rows for c in row.cells
        )
        assert "Power Board" in all_cell_text

    def test_multiple_assemblies_both_appear_in_table(self, tmp_path):
        """When multiple assemblies are present, all appear as section dividers."""
        results = [
            _make_result(_make_enriched("A", assembly="Board A"), [_make_cert()]),
            _make_result(_make_enriched("B", assembly="Board B"), [_make_cert()]),
        ]
        out = _generate(tmp_path, _make_pipeline_output(results))
        doc = Document(str(out))
        all_text = " ".join(c.text for row in doc.tables[0].rows for c in row.cells)
        assert "Board A" in all_text
        assert "Board B" in all_text


# ---------------------------------------------------------------------------
# Hyperlinks
# ---------------------------------------------------------------------------


class TestHyperlinks:
    def test_hyperlink_element_created_for_cert_with_url(self, tmp_path):
        """A w:hyperlink XML element is present when a cert has a source URL."""
        result = _make_result(
            certifications=[_make_cert(source_url="https://ul.com/listing/E171376")]
        )
        out = _generate(tmp_path, _make_pipeline_output([result]))
        doc = Document(str(out))
        hyperlinks = doc.element.body.findall(".//" + qn("w:hyperlink"))
        assert len(hyperlinks) >= 1

    def test_hyperlink_target_url_matches_source_url(self, tmp_path):
        """The hyperlink relationship target matches the cert source_url."""
        target_url = "https://productiq.ulprospector.com/en/profile/E171376"
        result = _make_result(
            certifications=[_make_cert(source_url=target_url, source_name="UL Product iQ")]
        )
        out = _generate(tmp_path, _make_pipeline_output([result]))
        doc = Document(str(out))
        hyperlinks = doc.element.body.findall(".//" + qn("w:hyperlink"))
        assert hyperlinks, "No hyperlink elements found"
        r_id = hyperlinks[0].get(qn("r:id"))
        assert r_id is not None
        rel = doc.part.rels.get(r_id)
        assert rel is not None
        assert rel.target_ref == target_url

    def test_no_hyperlink_when_source_url_absent(self, tmp_path):
        """No w:hyperlink element is added for a component with no source URL."""
        result = _make_result(certifications=[], confidence=Confidence.NOT_FOUND)
        out = _generate(tmp_path, _make_pipeline_output([result]))
        doc = Document(str(out))
        hyperlinks = doc.element.body.findall(".//" + qn("w:hyperlink"))
        assert len(hyperlinks) == 0

    def test_multiple_hyperlinks_for_multiple_certs(self, tmp_path):
        """One hyperlink per certification with a source URL."""
        certs = [
            _make_cert("UL 508", "https://ul.com/E1"),
            _make_cert("CE", "https://meanwell.com/CE"),
        ]
        result = _make_result(certifications=certs)
        out = _generate(tmp_path, _make_pipeline_output([result]))
        doc = Document(str(out))
        hyperlinks = doc.element.body.findall(".//" + qn("w:hyperlink"))
        assert len(hyperlinks) >= 2


# ---------------------------------------------------------------------------
# Footer note
# ---------------------------------------------------------------------------


class TestFooterNote:
    def test_footer_note_present_when_review_items_exist(self, tmp_path):
        """Italicised review note is added to the document body when needed."""
        data = _make_pipeline_output(
            components_needing_review=["Widget X (no certifications found)"]
        )
        out = _generate(tmp_path, data)
        doc = Document(str(out))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Widget X" in full_text
        assert "manual review" in full_text.lower()

    def test_footer_note_absent_when_no_review_items(self, tmp_path):
        """No review note paragraph is written when the review list is empty."""
        data = _make_pipeline_output(components_needing_review=[])
        out = _generate(tmp_path, data)
        doc = Document(str(out))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "manual review" not in full_text.lower()

    def test_footer_note_lists_all_flagged_components(self, tmp_path):
        """All items in components_needing_review appear in the footer note."""
        items = ["Component A ?", "Component B (no certifications found)"]
        data = _make_pipeline_output(components_needing_review=items)
        out = _generate(tmp_path, data)
        doc = Document(str(out))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Component A" in full_text
        assert "Component B" in full_text
